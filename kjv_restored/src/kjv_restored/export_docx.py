"""DOCX export for full Bible."""

from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, Iterator, Tuple, List
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DOCXExporter:
    """Exports Bible to DOCX format."""
    
    def __init__(self, title: str, version: str):
        """
        Initialize DOCX exporter.
        
        Args:
            title: Document title
            version: Version string
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        self.title = title
        self.version = version
        self.doc = Document()
        self._setup_document()
        self.current_page = 1
    
    def _setup_document(self):
        """Set up document styles and initial structure."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def add_title_page(self):
        """Add title page."""
        # Title
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(self.title)
        title_run.font.size = Pt(24)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spacing
        
        # Version
        version_para = self.doc.add_paragraph()
        version_run = version_para.add_run(self.version)
        version_run.font.size = Pt(14)
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spacing
        
        # Date
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Spacing
        self.doc.add_paragraph()  # Spacing
        
        # Disclaimer
        disclaimer_para = self.doc.add_paragraph()
        disclaimer_run = disclaimer_para.add_run(
            "Generated from user-supplied KJV text using restored-name rules. "
            "This document does not include or redistribute text from Cepher Bible or Dâbâr Yahuah."
        )
        disclaimer_run.font.size = Pt(10)
        disclaimer_run.italic = True
        disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        self.doc.add_page_break()
    
    def add_table_of_contents(self, books: List[str]):
        """
        Add table of contents.
        
        Args:
            books: List of book names in order
        """
        toc_para = self.doc.add_heading('Table of Contents', level=1)
        
        for book in books:
            para = self.doc.add_paragraph(book, style='List Bullet')
        
        self.doc.add_page_break()
    
    def add_book_heading(self, book_name: str):
        """
        Add book heading.
        
        Args:
            book_name: Name of the book
        """
        # Add section break for new book (if not first)
        if self.current_page > 1:
            # Add some spacing
            self.doc.add_paragraph()
        
        heading = self.doc.add_heading(book_name, level=1)
    
    def add_chapter_heading(self, chapter_num: int):
        """
        Add chapter heading.
        
        Args:
            chapter_num: Chapter number
        """
        heading = self.doc.add_heading(f"Chapter {chapter_num}", level=2)
    
    def add_verse(self, verse_num: int, text: str):
        """
        Add verse text.
        
        Args:
            verse_num: Verse number
            text: Verse text
        """
        para = self.doc.add_paragraph()
        
        # Add verse number as superscript
        verse_run = para.add_run(f"{verse_num}")
        verse_run.font.superscript = True
        verse_run.font.size = Pt(9)
        
        # Add space and text
        para.add_run(f" {text}")
    
    def add_footer(self):
        """Add footer with title and page number."""
        # Note: python-docx has limitations with page numbers
        # We'll add footer text, but true page numbers require more complex setup
        for section in self.doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"{self.title} - Page "
            # True page numbers would require field codes which are complex
    
    def export(self, events: Iterator[Tuple[str, Dict[str, Any]]], output_path: Path, progress_callback: Optional[callable] = None):
        """
        Export Bible from event stream.
        
        Args:
            events: Iterator of (event_type, event_data) tuples
            output_path: Output file path
            progress_callback: Optional progress callback
        """
        books_seen = []
        current_book = None
        current_chapter = None
        
        # Add title page
        self.add_title_page()
        
        # Process events
        for event_type, event_data in events:
            if event_type == "book":
                book_name = event_data["name"]
                if book_name not in books_seen:
                    books_seen.append(book_name)
                current_book = book_name
                self.add_book_heading(book_name)
                if progress_callback:
                    progress_callback(book_name, None, None)
            
            elif event_type == "chapter":
                chapter_num = event_data["number"]
                current_chapter = chapter_num
                self.add_chapter_heading(chapter_num)
            
            elif event_type == "verse":
                verse_num = event_data["verse"]
                text = event_data["text"]
                self.add_verse(verse_num, text)
                if progress_callback:
                    progress_callback(current_book, current_chapter, verse_num)
            
            elif event_type == "end":
                break
        
        # Add footer
        self.add_footer()
        
        # Save document
        self.doc.save(str(output_path))
    
    @staticmethod
    def is_available() -> bool:
        """Check if DOCX export is available."""
        return DOCX_AVAILABLE

